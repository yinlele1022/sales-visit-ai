"""
Sales Visit AI Analysis System - Web Product
============================================
Complete runnable product with:
1. Web interface for input (text/file upload)
2. AI analysis with fallback chain (deepseek -> qwen)
3. In-page card display with Confirm/Modify/Cancel buttons
4. Save to CSV table after confirmation
5. Audit logging
6. Feishu notification
7. Memory store for customer history

Run: python web_app.py
Open: http://localhost:5000
"""

import csv
import json
import os
import sys
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from flask import (
    Flask,
    render_template_string,
    request,
    jsonify,
    send_from_directory,
)

# Import project modules
from robust_agent import process_with_guardrails
from feishu_notifier import FeishuNotifier
from memory_store import MemoryStore
from feedback_loop import FeedbackStore
from audit_logger import ExcelAuditLogger, AuditEntry

# ============================================================================
# Configuration
# ============================================================================

APP_DIR = Path(__file__).parent
OUTPUT_DIR = APP_DIR / "outputs"
LOG_DIR = APP_DIR / "audit_logs"
DATA_DIR = APP_DIR / "data"

OUTPUT_FILE = OUTPUT_DIR / "Sales_Visit_Record_Table.csv"

# Ensure directories exist
for d in [OUTPUT_DIR, LOG_DIR, DATA_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# Initialize Flask app
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10MB max file size
app.secret_key = "sales_visit_ai_product_2026"

# Initialize modules
notifier = FeishuNotifier()
memory_store = MemoryStore()
feedback_store = FeedbackStore()

# In-memory session storage for pending records
pending_records: dict[str, dict] = {}

# ============================================================================
# HTML Templates (Embedded - All Chinese)
# ============================================================================

HOME_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>销售拜访 AI 分析系统</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Microsoft YaHei', sans-serif; background: #f0f2f5; color: #1a1a1a; min-height: 100vh; }
.header { background: linear-gradient(135deg, #1a56db 0%, #7c3aed 100%); color: white; padding: 24px 0; text-align: center; }
.header h1 { font-size: 24px; font-weight: 600; letter-spacing: -0.5px; }
.header p { font-size: 14px; opacity: 0.85; margin-top: 4px; }
.container { max-width: 720px; margin: 32px auto; padding: 0 20px; }

.card { background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.06); padding: 28px; margin-bottom: 20px; }
.card-title { font-size: 16px; font-weight: 600; color: #1a1a1a; margin-bottom: 16px; display: flex; align-items: center; gap: 8px; }
.card-title .icon { width: 22px; height: 22px; background: #eef2ff; border-radius: 6px; display: flex; align-items: center; justify-content: center; font-size: 13px; color: #1a56db; }

.form-group { margin-bottom: 18px; }
label { font-size: 13px; font-weight: 600; color: #374151; display: block; margin-bottom: 6px; }
input[type="text"], textarea, input[type="file"] { width: 100%; padding: 11px 14px; border: 1.5px solid #d1d5db; border-radius: 8px; font-size: 14px; transition: border-color 0.2s; background: #fafafa; }
textarea { min-height: 160px; resize: vertical; line-height: 1.6; font-family: inherit; }
input[type="text"]:focus, textarea:focus { outline: none; border-color: #1a56db; background: white; box-shadow: 0 0 0 3px rgba(26,86,219,0.08); }
.file-upload-area { border: 2px dashed #d1d5db; border-radius: 8px; padding: 28px; text-align: center; cursor: pointer; transition: all 0.2s; background: #fafafa; }
.file-upload-area:hover { border-color: #1a56db; background: #f0f4ff; }
.file-upload-area.has-file { border-style: solid; border-color: #10b981; background: #ecfdf5; }
.file-upload-area p { font-size: 13px; color: #6b7280; }
.file-name { color: #059669; font-weight: 500; }

.btn-row { display: flex; gap: 12px; margin-top: 20px; }
.btn { flex: 1; padding: 12px 20px; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; transition: all 0.15s; text-align: center; text-decoration: none; display: inline-block; }
.btn-primary { background: #1a56db; color: white; }
.btn-primary:hover { background: #1544b0; transform: translateY(-1px); }
.btn-primary:disabled { background: #9ca3af; cursor: not-allowed; transform: none; }
.btn-secondary { background: #f3f4f6; color: #374151; border: 1px solid #d1d5db; }
.btn-secondary:hover { background: #e5e7eb; }

.stats-row { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin-bottom: 20px; }
.stat-card { background: white; border-radius: 10px; padding: 16px; text-align: center; box-shadow: 0 1px 4px rgba(0,0,0,0.04); }
.stat-num { font-size: 24px; font-weight: 700; color: #1a56db; }
.stat-label { font-size: 12px; color: #6b7280; margin-top: 2px; }

.loading-overlay { display: none; position: fixed; inset: 0; background: rgba(255,255,255,0.88); z-index: 999; justify-content: center; align-items: center; flex-direction: column; gap: 20px; }
.loading-overlay.active { display: flex; }
.progress-container { width: 320px; height: 8px; background: #e5e7eb; border-radius: 4px; overflow: hidden; }
.progress-bar { height: 100%; background: linear-gradient(90deg, #1a56db, #3b82f6); border-radius: 4px; width: 0%; transition: width 0.4s ease; }
.loading-text { font-size: 14px; color: #374151; font-weight: 500; }

.error-banner { background: #fef2f2; border: 1px solid #fecaca; color: #dc2626; padding: 12px 16px; border-radius: 8px; font-size: 13px; margin-bottom: 16px; display: none; }
.error-banner.show { display: block; }

.footer { text-align: center; padding: 24px; color: #9ca3af; font-size: 12px; }
</style>
</head>
<body>

<div class="header">
<h1>销售拜访 AI 分析系统</h1>
<p>输入拜访记录 &raquo; AI 智能提取 &raquo; 人工确认 &raquo; 写入表格</p>
</div>

<div class="container">

<!-- Stats -->
<div class="stats-row">
<div class="stat-card"><div class="stat-num" id="total-count">0</div><div class="stat-label">已入库记录</div></div>
<div class="stat-card"><div class="stat-num" id="pending-count">0</div><div class="stat-label">待确认</div></div>
<div class="stat-card"><div class="stat-num" id="customer-count">0</div><div class="stat-label">客户数</div></div>
</div>

<!-- Error Banner -->
<div class="error-banner" id="error-banner"></div>

<!-- Input Form -->
<div class="card">
<div class="card-title"><span class="icon">1</span>新建拜访记录</div>

<form id="submit-form" method="post" action="/api/analyze" enctype="multipart/form-data">
<div class="form-group">
<label>销售姓名</label>
<input type="text" name="sales_name" placeholder="例如：张伟" required value="">
</div>

<div class="form-group">
<label>客户名称（选填）</label>
<input type="text" name="customer" placeholder="例如：长城汽车" value="">
</div>

<div class="form-group">
<label>拜访记录（粘贴文本或上传文件）</label>
<textarea name="transcript" placeholder="在此粘贴拜访记录，或拖拽文件到下方上传区域...&#10;&#10;示例：&#10;今天拜访了长城汽车，见到了王总（项目经理）。讨论了 Mpilot Highway 方案..." ></textarea>
</div>

<div class="form-group">
<label>或上传文件（.txt / .csv / .docx / .json）</label>
<div class="file-upload-area" id="file-upload-area">
<input type="file" name="file" accept=".txt,.csv,.docx,.json" id="file-input" style="display:none;">
<p id="file-upload-text">点击此处选择文件，或将文件拖拽到此处</p>
<p id="file-name-display" class="file-name" style="display:none;"></p>
</div>
</div>

<div class="btn-row">
<button type="submit" class="btn btn-primary" id="analyze-btn">AI 智能分析</button>
<button type="button" class="btn btn-secondary" id="view-history-btn">查看所有记录</button>
</div>
</form>
</div>

</div>

<div class="footer">销售拜访 AI 分析系统 v1.0 | 基于 Python + DeepSeek/Qwen LLM</div>

<!-- Loading Overlay -->
<div class="loading-overlay" id="loading">
<div class="progress-container"><div class="progress-bar" id="progress-bar"></div></div>
<div class="loading-text" id="loading-text">AI 分析中...</div>
</div>

<script>
// File upload - click to open file picker
const fileInput = document.getElementById('file-input');
const uploadArea = document.getElementById('file-upload-area');
const uploadText = document.getElementById('file-upload-text');
const fileNameDisplay = document.getElementById('file-name-display');

uploadArea.addEventListener('click', function() { fileInput.click(); });

fileInput.addEventListener('change', function(e) {
const file = e.target.files[0];
if (file) {
var ext = file.name.split('.').pop().toLowerCase();
if (ext === 'docx') {
// DOCX: don't preview binary, backend handles extraction
document.querySelector('textarea[name="transcript"]').value = '';
} else {
const reader = new FileReader();
reader.onload = function(ev) {
document.querySelector('textarea[name="transcript"]').value = ev.target.result;
};
reader.readAsText(file);
}
uploadText.style.display = 'none';
fileNameDisplay.style.display = 'block';
fileNameDisplay.textContent = ext === 'docx'
? '已选择: ' + file.name + '（Word 文档，将由后端解析）'
: '已选择: ' + file.name;
uploadArea.classList.add('has-file');
}
});

// Drag and drop support
uploadArea.addEventListener('dragover', function(e) { e.preventDefault(); uploadArea.style.borderColor = '#1a56db'; uploadArea.style.background = '#f0f4ff'; });
uploadArea.addEventListener('dragleave', function() { uploadArea.style.borderColor = ''; uploadArea.style.background = ''; });
uploadArea.addEventListener('drop', function(e) {
e.preventDefault();
uploadArea.style.borderColor = '';
uploadArea.style.background = '';
const file = e.dataTransfer.files[0];
if (file) {
fileInput.files = e.dataTransfer.files;
fileInput.dispatchEvent(new Event('change'));
}
});

// Form submission
document.getElementById('submit-form').addEventListener('submit', function(e) {
e.preventDefault();

// Frontend validation: need text OR file
var hasText = document.querySelector('textarea[name="transcript"]').value.trim();
var hasFile = document.getElementById('file-input').files.length > 0;
if (!hasText && !hasFile) {
var eb = document.getElementById('error-banner');
eb.textContent = '请粘贴拜访记录或上传文件';
eb.classList.add('show');
return;
}

const btn = document.getElementById('analyze-btn');
const overlay = document.getElementById('loading');
const loadingText = document.getElementById('loading-text');
const errorBanner = document.getElementById('error-banner');

// Show loading with progress bar
errorBanner.classList.remove('show');
overlay.classList.add('active');
btn.disabled = true;
var progressBar = document.getElementById('progress-bar');
progressBar.style.width = '0%';
loadingText.textContent = 'AI 分析中...';

var startTime = Date.now();
var timerInterval = setInterval(function() {
var elapsed = Math.floor((Date.now() - startTime) / 1000);
var fakeProgress = Math.min(90, elapsed * 12); // ~12% per second, cap at 90%
progressBar.style.width = fakeProgress + '%';
loadingText.textContent = 'AI 分析中（' + elapsed + ' 秒）...';
}, 500);

// Collect form data
const formData = new FormData(this);

fetch('/api/analyze', {
method: 'POST',
body: formData
})
.then(r => r.json())
.then(data => {
clearInterval(timerInterval);
progressBar.style.width = '100%';
if (data.success) {
if (data.batch_mode) {
// Batch processing result
setTimeout(function() { overlay.classList.remove('active'); }, 500);
btn.disabled = false;
var total = data.total;
var auto = data.auto_confirmed_count || 0;
var processed = data.processed || 0;
loadingText.textContent = '完成！共 ' + total + ' 条，' + auto + ' 条自动入库';
errorBanner.textContent = '批量处理完成：共 ' + total + ' 条，成功 ' + processed + ' 条（' + auto + ' 条自动入库）';
errorBanner.classList.add('show');
errorBanner.style.background = '#d1fae5';
errorBanner.style.borderColor = '#a7f3d0';
errorBanner.style.color = '#065f46';
// Refresh stats
fetch('/api/stats').then(r => r.json()).then(d => {
document.getElementById('total-count').textContent = d.total;
document.getElementById('pending-count').textContent = d.pending;
document.getElementById('customer-count').textContent = d.customers;
});
} else if (data.auto_confirmed) {
loadingText.textContent = '完成！已自动入库（置信度 ' + data.confidence + '）';
setTimeout(function() { overlay.classList.remove('active'); }, 500);
btn.disabled = false;
errorBanner.textContent = '记录已自动入库。飞书已发送通知，可在「查看所有记录」中修改。';
errorBanner.classList.add('show');
errorBanner.style.background = '#d1fae5';
errorBanner.style.borderColor = '#a7f3d0';
errorBanner.style.color = '#065f46';
} else {
loadingText.textContent = '完成，跳转到确认页面...';
setTimeout(function() { window.location.href = '/review/' + data.record_id; }, 400);
}
} else {
errorBanner.textContent = data.error || '分析失败，请重试';
errorBanner.classList.add('show');
overlay.classList.remove('active');
btn.disabled = false;
}
})
.catch(err => {
errorBanner.textContent = '网络错误：' + err.message;
errorBanner.classList.add('show');
overlay.classList.remove('active');
btn.disabled = false;
});
});

// Load stats on page load
fetch('/api/stats').then(r => r.json()).then(d => {
document.getElementById('total-count').textContent = d.total;
document.getElementById('pending-count').textContent = d.pending;
document.getElementById('customer-count').textContent = d.customers;
});

// History button - navigate outside form
document.getElementById('view-history-btn').addEventListener('click', function() {
window.location.href = '/history';
});
</script>
</body>
</html>
"""

REVIEW_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>确认记录 - 销售拜访 AI</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Microsoft YaHei', sans-serif; background: #f0f2f5; color: #1a1a1a; min-height: 100vh; }
.header { background: linear-gradient(135deg, #1a56db 0%, #7c3aed 100%); color: white; padding: 20px 0; }
.header-inner { max-width: 720px; margin: 0 auto; padding: 0 20px; display: flex; align-items: center; gap: 12px; }
.back-btn { color: rgba(255,255,255,0.8); text-decoration: none; font-size: 14px; display: flex; align-items: center; gap: 4px; }
.back-btn:hover { color: white; }
.header h1 { font-size: 18px; font-weight: 600; }
.container { max-width: 720px; margin: 24px auto; padding: 0 20px; }
.status-badge { display: inline-flex; align-items: center; gap: 6px; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; }
.status-pending { background: #fef3c7; color: #92400e; }
.status-saved { background: #d1fae5; color: #065f46; }
.card { background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.06); padding: 24px; margin-bottom: 16px; overflow: hidden; }
.card-header { display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px; }
.card-header h2 { font-size: 17px; font-weight: 600; }

/* Result card - mimics Feishu card style */
.result-card { border-left: 4px solid #1a56db; }
.field-grid { display: grid; grid-template-columns: 120px 1fr; gap: 8px 16px; font-size: 14px; }
.field-label { color: #6b7280; font-size: 13px; font-weight: 500; padding-top: 2px; }
.field-value { color: #1a1a1a; font-weight: 500; word-break: break-word; }
.field-value.risk-high { color: #dc2626; font-weight: 700; }
.field-value.risk-mid { color: #d97706; font-weight: 700; }
.field-value.risk-low { color: #059669; font-weight: 700; }

.section-divider { height: 1px; background: #e5e7eb; margin: 16px 0; }
.section-title { font-size: 12px; font-weight: 700; color: #6b7280; letter-spacing: 0.5px; margin-bottom: 8px; }

.tag-list { display: flex; flex-wrap: wrap; gap: 6px; }
.tag { background: #eff6ff; color: #1d4ed8; padding: 3px 10px; border-radius: 4px; font-size: 12px; font-weight: 500; }
.tag.concern { background: #fef3c7; color: #92400e; }
.tag.competitor { background: #fce7f3; color: #be185d; }

.action-buttons { display: flex; gap: 10px; margin-top: 20px; }
.btn { padding: 11px 24px; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; transition: all 0.15s; flex: 1; text-align: center; }
.btn-confirm { background: #059669; color: white; }
.btn-confirm:hover { background: #047857; }
.btn-modify { background: #f59e0b; color: white; }
.btn-modify:hover { background: #d97706; }
.btn-cancel { background: #ef4444; color: white; }
.btn-cancel:hover { background: #dc2626; }
.btn:disabled { opacity: 0.5; cursor: not-allowed; }

.success-message { background: #d1fae5; border: 1px solid #a7f3d0; color: #065f46; padding: 16px; border-radius: 8px; text-align: center; margin-bottom: 16px; display: none; }
.success-message.show { display: block; }

.original-text { background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 8px; padding: 14px; font-size: 13px; color: #4b5563; line-height: 1.6; max-height: 150px; overflow-y: auto; }

.modify-panel { display: none; background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); padding: 24px; margin-bottom: 16px; border: 2px solid #f59e0b; }
.modify-panel.show { display: block; }
.modify-field { margin-bottom: 12px; }
.modify-field label { font-size: 12px; font-weight: 600; color: #374151; display: block; margin-bottom: 4px; }
.modify-field input, .modify-field select, .modify-field textarea { width: 100%; padding: 8px 12px; border: 1.5px solid #d1d5db; border-radius: 6px; font-size: 13px; font-family: inherit; }

.info-row { display: flex; gap: 16px; font-size: 12px; color: #6b7280; margin-bottom: 16px; }
.info-item { display: flex; align-items: center; gap: 4px; }
</style>
</head>
<body>

<div class="header">
<div class="header-inner">
<a href="/" class="back-btn">&larr; 返回</a>
<h1>确认 & 入库</h1>
<span class="status-badge status-pending" id="status-badge">待确认</span>
</div>
</div>

<div class="container">

<!-- Success Message -->
<div class="success-message" id="success-msg">
<strong>入库成功！</strong> 记录已写入表格。
<br><br>
<a href="/" style="color: #059669;">&larr; 新建记录</a>&nbsp;&nbsp;&nbsp;
<a href="/history" style="color: #059669;">查看所有记录 &rarr;</a>
</div>

<!-- Result Card -->
<div class="card result-card" id="result-card">
<div class="card-header">
<h2>AI 提取信息</h2>
<span id="model-info" style="font-size:12px;color:#9ca3af;"></span>
</div>

<div class="field-grid">
<div class="field-label">销售</div><div class="field-value" id="val-sales"></div>
<div class="field-label">客户</div><div class="field-value" id="val-customer"></div>
<div class="field-label">联系人</div><div class="field-value" id="val-contact"></div>
<div class="field-label">拜访时间</div><div class="field-value" id="val-time"></div>
<div class="field-label">核心议题</div><div class="field-value" id="val-topic"></div>
</div>

<div class="section-divider"></div>
<div class="section-title">风险评估</div>
<div class="field-grid">
<div class="field-label">风险等级</div><div class="field-value" id="val-risk"></div>
<div class="field-label">风险原因</div><div class="field-value" id="val-risk-reason"></div>
</div>

<div class="section-divider"></div>
<div class="section-title">关键疑虑</div>
<div class="tag-list" id="val-concerns"><span style="color:#9ca3af;font-size:13px;">无</span></div>

<div class="section-divider"></div>
<div class="section-title">竞对动态</div>
<div class="tag-list" id="val-competitors"><span style="color:#9ca3af;font-size:13px;">无</span></div>

<div class="section-divider"></div>
<div class="section-title">下一步行动</div>
<div id="val-next-steps" style="font-size:13px;color:#374151;line-height:1.8;"></div>

<div class="section-divider"></div>
<div class="section-title">战略洞察</div>
<div id="val-strategic-insight" style="font-size:13px;color:#1a56db;line-height:1.8;background:#eff6ff;padding:10px 14px;border-radius:6px;border-left:3px solid #1a56db;"></div>

<!-- Action Buttons -->
<div class="action-buttons">
<button class="btn btn-confirm" id="btn-confirm" onclick="confirmRecord()">确认入库</button>
<button class="btn btn-modify" id="btn-modify" onclick="showModifyPanel()">修改</button>
<button class="btn btn-cancel" id="btn-cancel" onclick="cancelRecord()">取消</button>
</div>
</div>

<!-- Modify Panel -->
<div class="modify-panel" id="modify-panel">
<h3 style="margin-bottom:16px;font-size:15px;">修改 AI 提取结果</h3>
<div id="modify-fields"></div>
<div style="display:flex;gap:10px;margin-top:16px;">
<button class="btn btn-confirm" onclick="saveModification()">保存修改并入库</button>
<button class="btn btn-cancel" onclick="hideModifyPanel()">放弃修改</button>
</div>
</div>

<!-- Original Text -->
<div class="card">
<h3 style="font-size:14px;margin-bottom:10px;color:#6b7280;">原始输入文本</h3>
<div class="original-text" id="original-text"></div>
</div>

<!-- Processing Info -->
<div class="info-row">
<div class="info-item" id="info-model"></div>
<div class="info-item" id="info-time"></div>
<div class="info-item" id="info-fallback"></div>
</div>

</div>

<script>
const RECORD_ID = "{{ record_id }}";

// Load record data on page load
async function loadRecord() {
try {
const resp = await fetch('/api/record/' + RECORD_ID);
const data = await resp.json();

if (!data.success) {
alert('记录不存在或已过期');
window.location.href = '/';
return;
}

const ext = data.extracted;

// Fill in values
document.getElementById('val-sales').textContent = ext.sales_name || '-';
document.getElementById('val-customer').textContent = ext.customer || '-';
document.getElementById('val-contact').textContent = ext.contact_person || '-';
document.getElementById('val-time').textContent = ext.visit_time || '-';
document.getElementById('val-topic').textContent = ext.topic || '-';

// Risk level with color
const riskEl = document.getElementById('val-risk');
const riskVal = ext.risk_level || '-';
riskEl.textContent = riskVal;
riskEl.className = 'field-value ' + (riskVal.includes('高') ? 'risk-high' : riskVal.includes('中') ? 'risk-mid' : riskVal.includes('低') ? 'risk-low' : '');
document.getElementById('val-risk-reason').textContent = ext.risk_reason || '-';

// Concerns
renderTags('val-concerns', ext.key_concerns, 'concern');

// Competitors - support both field names
var compData = ext.competitors_mentioned || ext.competitors || [];
renderTags('val-competitors', compData, 'competitor');

// Next steps
renderNextSteps(ext.next_steps);

// Original text
document.getElementById('original-text').textContent = data.transcript || '';

// Info
document.getElementById('info-model').innerHTML = '<strong>模型：</strong>' + (data.provider_used || '-');
document.getElementById('info-time').innerHTML = '<strong>耗时：</strong>' + ((data.elapsed_seconds||0).toFixed(1)) + 's';
document.getElementById('info-fallback').innerHTML = '<strong>降级：</strong>' + (data.fallback_used ? '是' : '否');
document.getElementById('model-info').textContent = data.provider_used || '';
} catch (err) {
console.error(err);
}
}

function renderTags(elementId, items, className) {
const el = document.getElementById(elementId);
if (!items || !items.length || (Array.isArray(items) && items.length === 0)) return;
const arr = Array.isArray(items) ? items : [];
el.innerHTML = arr.map(function(item) {
let text = item;
if (typeof item === 'object') text = item.action || item.content || JSON.stringify(item);
return '<span class="tag ' + className + '">' + escapeHtml(text) + '</span>';
}).join('');
}

function renderNextSteps(items) {
const el = document.getElementById('val-next-steps');
if (!items || !items.length || (Array.isArray(items) && items.length === 0)) { el.innerHTML = '<span style="color:#9ca3af">无</span>'; return; }
const arr = Array.isArray(items) ? items : [];
el.innerHTML = arr.map(function(item) {
let text, deadline = '', owner = '';
if (typeof item === 'object') {
text = item.action || item.content || '';
deadline = item.deadline || '';
owner = item.owner || '';
} else {
text = String(item);
}
let line = '&bull; ' + escapeHtml(text);
if (deadline) line += ' <small style="color:#6b7280">（截止：' + escapeHtml(deadline) + '）</small>';
if (owner) line += ' <small style="color:#9ca3af">[' + escapeHtml(owner) + ']</small>';
return '<div style="margin-bottom:4px;">' + line + '</div>';
}).join('');
}

// Strategic insight
var insightEl = document.getElementById('val-strategic-insight');
var insight = ext.strategic_insight;
if (insight && insight !== '无' && insight !== '') {
insightEl.innerHTML = '💡 ' + escapeHtml(insight);
} else {
insightEl.innerHTML = '<span style="color:#9ca3af">无</span>';
}
}

function escapeHtml(str) {
if (!str) return '';
return String(str).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

// Confirm record
async function confirmRecord() {
await actionRecord('confirm');
}

// Cancel record
function cancelRecord() {
if (confirm('确定要取消这条记录吗？')) {
actionRecord('cancel');
}
}

// Show modify panel
function showModifyPanel() {
const ext = {{ extracted_json | safe }};
var compVal = ext.competitors_mentioned || ext.competitors || [];
const fields = [
{ key: 'sales_name', label: '销售姓名', value: ext.sales_name },
{ key: 'customer', label: '客户名称', value: ext.customer },
{ key: 'contact_person', label: '联系人', value: ext.contact_person },
{ key: 'visit_time', label: '拜访时间', value: ext.visit_time },
{ key: 'topic', label: '核心议题', value: ext.topic },
{ key: 'risk_level', label: '风险等级', value: ext.risk_level, options: ['高', '中', '低'] },
{ key: 'key_concerns', label: '关键疑虑（逗号分隔）', value: Array.isArray(ext.key_concerns) ? ext.key_concerns.join(', ') : ext.key_concerns },
{ key: 'competitors_mentioned', label: '竞对动态（逗号分隔）', value: Array.isArray(compVal) ? compVal.join(', ') : compVal },
{ key: 'next_steps', label: '下一步行动（每行一条）', value: Array.isArray(ext.next_steps) ? ext.next_steps.map(function(s) { return typeof s === 'object' ? s.action : s; }).join('\\n') : ext.next_steps },
];

const container = document.getElementById('modify-fields');
container.innerHTML = fields.map(function(f) {
return '<div class="modify-field">' +
'<label>' + f.label + '</label>' +
(f.options
? '<select name="' + f.key + '" id="mod-' + f.key + '">' + f.options.map(function(o) { return '<option' + (o===f.value?' selected':'') + '>' + o + '</option>'; }).join('') + '</select>'
: f.key === 'next_steps'
? '<textarea name="' + f.key + '" id="mod-' + f.key + '" rows="3">' + escapeHtml(f.value) + '</textarea>'
: '<input type="text" name="' + f.key + '" id="mod-' + f.key + '" value="' + escapeHtml(f.value) + '">') +
'</div>';
}).join('');

document.getElementById('modify-panel').classList.add('show');
document.getElementById('btn-modify').disabled = true;
document.getElementById('btn-confirm').disabled = true;
document.getElementById('btn-cancel').disabled = true;
}

function hideModifyPanel() {
document.getElementById('modify-panel').classList.remove('show');
document.getElementById('btn-modify').disabled = false;
document.getElementById('btn-confirm').disabled = false;
document.getElementById('btn-cancel').disabled = false;
}

// Save modification and confirm
async function saveModification() {
const modifiedData = {};
['sales_name','customer','contact_person','visit_time','topic','risk_level','key_concerns','competitors_mentioned','next_steps'].forEach(function(key) {
const el = document.getElementById('mod-' + key);
if (!el) return;
let val = el.value.trim();
if (key === 'key_concerns' || key === 'competitors_mentioned') {
modifiedData[key] = val.split(',').map(function(s){return s.trim();}).filter(Boolean);
} else if (key === 'next_steps') {
modifiedData[key] = val.split('\\n').map(function(s){return s.trim();}).filter(Boolean);
} else {
modifiedData[key] = val;
}
});

await actionRecord('modify', modifiedData);
}

// Send action to server
async function actionRecord(action, modifications) {
try {
const body = { action: action };
if (modifications) body.modifications = modifications;

const resp = await fetch('/api/record/{{ record_id }}/action', {
method: 'POST',
headers: {'Content-Type': 'application/json'},
body: JSON.stringify(body)
});

const result = await resp.json();
if (result.success) {
document.getElementById('status-badge').className = 'status-badge status-saved';
document.getElementById('status-badge').textContent = '已入库';
document.getElementById('success-msg').classList.add('show');
document.getElementById('result-card').style.opacity = '0.6';
document.getElementById('modify-panel').style.display = 'none';

// Disable buttons
document.querySelectorAll('.btn').forEach(function(b) { b.disabled = true; });
} else {
alert('错误：' + (result.error || '未知错误'));
}
} catch (err) {
alert('网络错误：' + err.message);
}
}

// Init
loadRecord();
</script>
</body>
</html>
"""

HISTORY_PAGE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>所有记录 - 销售拜访 AI</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Microsoft YaHei', sans-serif; background: #f0f2f5; color: #1a1a1a; min-height: 100vh; }
.header { background: linear-gradient(135deg, #1a56db 0%, #7c3aed 100%); color: white; padding: 20px 0; }
.header-inner { max-width: 960px; margin: 0 auto; padding: 0 20px; display: flex; align-items: center; gap: 12px; }
.back-btn { color: rgba(255,255,255,0.8); text-decoration: none; font-size: 14px; }
.back-btn:hover { color: white; }
.header h1 { font-size: 18px; font-weight: 600; }
.container { max-width: 960px; margin: 24px auto; padding: 0 20px; }
.table-card { background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,0.06); overflow: hidden; }
table { width: 100%; border-collapse: collapse; font-size: 13px; }
th { background: #f9fafb; padding: 12px 14px; text-align: left; font-weight: 600; color: #374151; font-size: 12px; letter-spacing: 0.3px; border-bottom: 1px solid #e5e7eb; white-space: nowrap; }
td { padding: 10px 14px; border-bottom: 1px solid #f3f4f6; vertical-align: top; color: #4b5563; }
tr:hover td { background: #f9fafb; }
tr:last-child td { border-bottom: none; }
.empty-state { text-align: center; padding: 48px 20px; color: #9ca3af; }
.empty-state h2 { font-size: 16px; color: #6b7280; margin-bottom: 8px; }
.badge { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 11px; font-weight: 600; }
.badge-high { background: #fee2e2; color: #dc2626; }
.badge-mid { background: #fef3c7; color: #d97706; }
.badge-low { background: #d1fae5; color: #059669; }
.badge- { background: #f3f4f6; color: #6b7280; }
.footer { text-align: center; padding: 24px; color: #9ca3af; font-size: 12px; }
</style>
</head>
<body>

<div class="header">
<div class="header-inner">
<a href="/" class="back-btn">&larr; 首页</a>
<h1>销售拜访记录总表</h1>
</div>
</div>

<div class="container">
<div class="table-card">
{% if records %}
<table>
<thead>
<tr>
<th>#</th><th>销售</th><th>客户</th><th>联系人</th><th>拜访时间</th><th>核心议题</th><th>风险</th><th>状态</th><th>入库时间</th>
</tr>
</thead>
<tbody>
{% for r in records %}
<tr>
<td>{{ loop.index }}</td>
<td>{{ r.Sales or r.sales_name or '-' }}</td>
<td>{{ r.Customer or r.customer or '-' }}</td>
<td>{{ r.Contact or r.contact_person or '-' }}</td>
<td>{{ r.VisitTime or r.visit_time or '-' }}</td>
<td>{{ (r.Topic or r.topic or '-')[:30] }}{{ '...' if (r.Topic or r.topic or '')|length > 30 else '' }}</td>
<td><span class="badge badge-{{ (r.RiskLevel or r.risk_level or '')|lower }}">{{ r.RiskLevel or r.risk_level or '-' }}</span></td>
<td>{{ r.ProcessResult or r.process_result or '-' }}</td>
<td>{{ r.SavedAt or r.saved_at or '-' }}</td>
</tr>
{% endfor %}
</tbody>
</table>
{% else %}
<div class="empty-state">
<h2>暂无记录</h2>
<p>从<a href="/">首页</a>提交您的第一条拜访记录。</p>
</div>
{% endif %}
</div>
</div>

<div class="footer">共 {{ total_count }} 条记录</div>
</body>
</html>
"""


# ============================================================================
# Helper Functions
# ============================================================================


def _get_next_id() -> int:
    """Get next record ID from CSV."""
    if not OUTPUT_FILE.exists():
        return 1
    try:
        with open(OUTPUT_FILE, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows = list(reader)
            if len(rows) <= 1:
                return 1
            last_id = int(rows[-1][0]) if rows[-1][0].isdigit() else 0
            return last_id + 1
    except Exception:
        return 1


def _save_to_csv(extracted: dict, transcript: str = "") -> int:
    """Save extracted record to CSV. Returns the row ID."""
    next_id = _get_next_id()

    # Ensure header exists
    if not OUTPUT_FILE.exists():
        with open(OUTPUT_FILE, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f, quoting=csv.QUOTE_ALL)
            writer.writerow([
                "ID", "Sales", "Customer", "Contact", "VisitTime",
                "Topic", "Interest", "Concerns", "Competitors",
                "NextSteps", "FollowUp", "RiskLevel", "RiskReason",
                "Resources", "RelationshipTip", "StrategicInsight", "ProcessResult",
                "SavedAt", "OriginalText"
            ])

    def fmt(val):
        """Format any value for CSV cell."""
        if isinstance(val, list):
            parts = []
            for item in val:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, dict):
                    parts.append(item.get("action") or item.get("content") or str(item))
                else:
                    parts.append(str(item))
            return "; ".join(parts)
        return str(val) if val else ""

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    row = [
        next_id,
        fmt(extracted.get("sales_name")),
        fmt(extracted.get("customer")),
        fmt(extracted.get("contact_person")),
        fmt(extracted.get("visit_time")),
        fmt(extracted.get("topic")),
        fmt(extracted.get("customer_interest")),
        fmt(extracted.get("key_concerns")),
        # Support both competitors and competitors_mentioned
        fmt(extracted.get("competitors_mentioned") or extracted.get("competitors")),
        fmt(extracted.get("next_steps")),
        "",  # Follow-up time placeholder
        fmt(extracted.get("risk_level")),
        fmt(extracted.get("risk_reason")),
        fmt(extracted.get("resources_needed")),
        fmt(extracted.get("relationship_suggestion")),
        fmt(extracted.get("strategic_insight")),
        "已确认入库",
        now,
        (transcript or "")[:300],
    ]

    with open(OUTPUT_FILE, "a", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, quoting=csv.QUOTE_ALL)
        writer.writerow(row)

    return next_id


def _load_all_records() -> list[dict]:
    """Load all saved records from CSV."""
    if not OUTPUT_FILE.exists():
        return []

    records = []
    try:
        with open(OUTPUT_FILE, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                records.append(row)
    except Exception:
        pass
    return records


# ============================================================================
# Routes
# ============================================================================


@app.route("/")
def home():
    """Homepage - input form."""
    return render_template_string(HOME_PAGE)


@app.route("/history")
def history():
    """View all saved records."""
    records = _load_all_records()
    return render_template_string(HISTORY_PAGE, records=records, total_count=len(records))


@app.route("/review/<record_id>")
def review(record_id: str):
    """Review page - display card with confirm/modify/cancel buttons."""
    if record_id not in pending_records:
        return "记录不存在或已过期。<a href='/'>返回首页</a>", 404

    rec = pending_records[record_id]
    extracted_json = json.dumps(rec["extracted"], ensure_ascii=False, indent=2)

    return render_template_string(
        REVIEW_PAGE,
        record_id=record_id,
        extracted_json=extracted_json,
    )


@app.route("/outputs/<filename>")
def serve_output(filename: str):
    """Serve output files for download."""
    return send_from_directory(str(OUTPUT_DIR), filename)


# ============================================================================
# API Endpoints
# ============================================================================


@app.route("/api/analyze", methods=["POST"])
def api_analyze():
    """Analyze sales visit record with AI."""
    try:
        # Get form data
        sales_name = request.form.get("sales_name", "").strip()
        customer = request.form.get("customer", "").strip()
        transcript = request.form.get("transcript", "").strip()

        # Handle file upload - smart parsing for JSON/DOCX files
        uploaded_file = request.files.get("file")
        batch_mode = False
        batch_records = []
        import time as _time

        if uploaded_file and uploaded_file.filename:
            t0 = _time.time()
            raw_bytes = uploaded_file.read()
            fname_lower = uploaded_file.filename.lower()
            print(f"[Timing] File read: {_time.time()-t0:.2f}s, size={len(raw_bytes)} bytes")

            # DOCX/DOC: extract text from Word document
            if fname_lower.endswith(".docx") or fname_lower.endswith(".doc"):
                from io import BytesIO
                docx_text = None
                # Try python-docx first (for .docx / Office Open XML)
                try:
                    from docx import Document as DocxDocument
                    t1 = _time.time()
                    doc = DocxDocument(BytesIO(raw_bytes))
                    print(f"[Timing] docx parse: {_time.time()-t1:.2f}s, paragraphs={len(doc.paragraphs)}")
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    docx_text = "\n".join(paragraphs)
                    print(f"[Timing] docx extract: {_time.time()-t1:.2f}s, text_len={len(docx_text)}")
                except Exception as e1:
                    # Not a valid .docx — try .doc (OLE format, WPS/Word 97-2003)
                    try:
                        import olefile
                        ole = olefile.OleFileIO(BytesIO(raw_bytes))
                        # Read WordDocument stream
                        wd = ole.openstream("WordDocument").read()
                        # .doc stores text as UTF-16LE with binary markers interspersed
                        # Strategy: decode all, extract Chinese + punctuation sequences
                        text_16 = wd.decode("utf-16-le", errors="ignore")
                        # Extract readable text sequences (Chinese chars + common punctuation)
                        import re as _re
                        pattern = (
                            r'[\u4e00-\u9fff'        # CJK Unified Ideographs
                            r'\u3000-\u303f'          # CJK Symbols
                            r'\uff00-\uffef'          # Fullwidth forms
                            r'\w\s'                    # ASCII alphanumeric + whitespace
                            r'\u2018\u2019\u201c\u201d' # Smart quotes
                            r'\uff0c\uff0e\uff1a\uff1b\uff01\uff1f' # Fullwidth punctuation
                            r'\u3001\u3002'            # CJK punctuation
                            r'\.\,\!\?\;\:\(\)\[\]'   # ASCII punctuation
                            r'\#\*\-\+\@\/\d'          # Numbers and symbols
                            r']+'
                        )
                        chunks = _re.findall(pattern, text_16)
                        paragraphs = [c.strip() for c in chunks if len(c.strip()) > 10]
                        docx_text = "\n".join(paragraphs)
                        ole.close()
                        print(f"[Doc-OLD] Extracted {len(paragraphs)} paragraphs, {len(docx_text)} chars")
                    except Exception as e2:
                        print(f"[Doc] Both docx ({e1}) and ole ({e2}) failed, using raw decode")
                        docx_text = raw_bytes.decode("utf-8", errors="replace")

                if docx_text and not transcript:
                    transcript = docx_text
                    # If docx has many paragraphs (>=3), try to split into individual records
                    if len(paragraphs) >= 3:
                        # Each non-empty paragraph is treated as a separate visit record
                        records = [p for p in paragraphs if len(p) > 20]
                        if len(records) > 1:
                            batch_records = [
                                {"sales_name": sales_name, "customer": customer, "transcript": r}
                                for r in records
                            ]
                            batch_mode = True
                            print(f"[Docx] Split into {len(batch_records)} individual records for batch processing")

            # JSON: smart parsing
            elif fname_lower.endswith(".json"):
                raw = raw_bytes.decode("utf-8")
                try:
                    parsed = json.loads(raw)
                    if isinstance(parsed, list):
                        batch_records = [
                            {
                                "sales_name": r.get("sales_name", sales_name),
                                "customer": r.get("customer", customer),
                                "transcript": str(r.get("transcript", "")),
                            }
                            for r in parsed
                            if r.get("transcript")
                        ]
                        if batch_records:
                            batch_mode = True
                    elif isinstance(parsed, dict):
                        if not transcript:
                            transcript = str(parsed.get("transcript", raw))
                        if not sales_name:
                            sales_name = str(parsed.get("sales_name", "")).strip()
                        if not customer:
                            customer = str(parsed.get("customer", "")).strip()
                except Exception:
                    if not transcript:
                        transcript = raw

            # Plain text / CSV: read as UTF-8
            elif not transcript:
                transcript = raw_bytes.decode("utf-8", errors="replace")

        if not batch_mode and not transcript:
            return jsonify({"success": False, "error": "请提供拜访记录文本"}), 400
        if not batch_mode and not sales_name:
            return jsonify({"success": False, "error": "请填写销售姓名"}), 400

        # Batch mode: process all records
        if batch_mode:
            results = []
            for i, rec in enumerate(batch_records):
                try:
                    sub_transcript = rec["transcript"]
                    sub_sales = rec["sales_name"] or sales_name or "未知"
                    sub_customer = rec["customer"] or customer or "未知客户"

                    if not sub_transcript.strip():
                        continue

                    ctx = {"sales_name": sub_sales, "customer": sub_customer}
                    hist = ""
                    try:
                        hist = memory_store.get_context(sub_customer)
                    except Exception:
                        pass

                    kwargs = {
                        "transcript": sub_transcript,
                        "context": ctx,
                        "providers": ["deepseek", "qwen"],
                    }
                    if hist:
                        kwargs["history_context"] = hist

                    result = process_with_guardrails(**kwargs)
                    if result and result.extracted:
                        rid = str(uuid.uuid4())[:8]
                        ext = result.extracted
                        field_conf = result.confidence or {}
                        # Only average extraction-reliant fields for batch mode
                        extract_fields = ["contact_person","visit_time","topic","key_concerns",
                                         "competitors_mentioned","next_steps","risk_level","risk_reason"]
                        extract_scores = [float(v) for f in extract_fields for k,v in field_conf.items() 
                                         if k == f and float(v) > 0]
                        if not extract_scores:
                            extract_scores = [float(v) for v in field_conf.values() if float(v) > 0]
                        overall_conf = sum(extract_scores) / len(extract_scores) if extract_scores else 0
                        has_substance = len(sub_transcript) >= 25

                        rec_data = {
                            "sales_name": sub_sales,
                            "customer": sub_customer,
                            "transcript": sub_transcript,
                            "context": ctx,
                            "extracted": ext,
                            "provider_used": result.provider_used,
                            "fallback_used": result.fallback_used,
                            "elapsed_seconds": result.elapsed_seconds,
                            "confidence": result.confidence,
                            "passed": result.passed,
                            "confirmed": False,
                        }

                        if overall_conf > 0.85 and result.passed and has_substance:
                            rec_data["confirmed"] = True
                            try:
                                _handle_confirm(rec_data, rid)
                            except Exception:
                                rec_data["confirmed"] = False

                        pending_records[rid] = rec_data

                        results.append({
                            "index": i + 1,
                            "sales_name": sub_sales,
                            "customer": sub_customer,
                            "auto_confirmed": rec_data["confirmed"],
                            "confidence": round(overall_conf, 2),
                        })

                except Exception as e:
                    results.append({
                        "index": i + 1,
                        "error": str(e),
                    })

            auto_count = sum(1 for r in results if r.get("auto_confirmed"))
            return jsonify({
                "success": True,
                "batch_mode": True,
                "total": len(batch_records),
                "processed": len([r for r in results if "auto_confirmed" in r or "error" in r]),
                "auto_confirmed_count": auto_count,
                "results": results[:20],  # Cap at 20 for response size
            })

        # Build context
        context = {"sales_name": sales_name, "customer": customer or "未知客户"}

        # Truncate long transcripts to speed up API calls (keep full text for storage)
        MAX_TRANSCRIPT_LEN = 3000
        full_transcript = transcript
        if len(transcript) > MAX_TRANSCRIPT_LEN:
            transcript = transcript[:MAX_TRANSCRIPT_LEN] + "\n...(内容过长已截断)"
            print(f"[Truncate] transcript {len(full_transcript)} -> {len(transcript)} chars")

        # Get customer history from memory store
        history_context = ""
        try:
            hist = memory_store.get_context(customer)
            if hist:
                history_context = hist
        except Exception:
            pass

        # Run AI analysis - use fallback chain: deepseek first, then qwen
        print(f"[Timing] Starting AI analysis, transcript_len={len(transcript)}")
        t_ai = _time.time()
        kwargs = {
            "transcript": transcript,
            "context": context,
            "providers": ["deepseek", "qwen"],
        }
        if history_context:
            kwargs["history_context"] = history_context

        result = process_with_guardrails(**kwargs)
        print(f"[Timing] AI analysis done: {_time.time()-t_ai:.1f}s")

        if not result or not result.extracted:
            error_msg = "AI 分析未返回结果"
            if result and result.human_reason:
                error_msg += f"：{result.human_reason}"
            return jsonify({"success": False, "error": error_msg}), 500

        # Convert extracted dict to our format
        # result.extracted is a dict (from model_dump()), use .get() not getattr()
        extracted = {}
        field_mapping = {
            "sales_name": "sales_name",
            "customer": "customer",
            "contact_person": "contact_person",
            "visit_time": "visit_time",
            "topic": "topic",
            "customer_interest": "customer_interest",
            "key_concerns": "key_concerns",
            "competitors_mentioned": "competitors_mentioned",
            "next_steps": "next_steps",
            "risk_level": "risk_level",
            "risk_reason": "risk_reason",
            "resources_needed": "resources_needed",
            "relationship_suggestion": "relationship_suggestion",
            "strategic_insight": "strategic_insight",
        }
        for schema_field, our_field in field_mapping.items():
            val = result.extracted.get(schema_field)
            if val is not None:
                extracted[our_field] = val

        # Fix common LLM extraction errors: fields sometimes get swapped
        import re as _re

        def _is_likely_date(val):
            return bool(val and _re.match(r'^\d{4}-\d{2}-\d{2}', str(val)))

        def _is_likely_topic(val):
            """True if val looks like a topic (Chinese text, not a date)"""
            s = str(val).strip()
            return bool(s and s not in ("未知","null","None","") and not _is_likely_date(s)
                       and _re.search(r'[\u4e00-\u9fff]', s))

        topic_val = str(extracted.get("topic", ""))
        visit_val = str(extracted.get("visit_time", ""))

        # Case 1: topic contains a date prefix → move to visit_time
        date_match = _re.match(r'^(\d{4}-\d{2}-\d{2})', topic_val)
        if date_match and not _is_likely_date(visit_val):
            extracted["visit_time"] = date_match.group(1)
            cleaned = topic_val[date_match.end():].strip()
            extracted["topic"] = cleaned if cleaned else "未知"

        # Case 2: visit_time is actually the topic (LLM swapped them)
        if _is_likely_topic(visit_val) and not _is_likely_topic(topic_val):
            extracted["topic"] = visit_val
            extracted["visit_time"] = datetime.now().strftime("%Y-%m-%d")

        # Case 3: visit_time still empty → fill today
        final_visit = str(extracted.get("visit_time", ""))
        if not final_visit or final_visit in ("未知", "null", "None", ""):
            extracted["visit_time"] = datetime.now().strftime("%Y-%m-%d")

        # Generate unique record ID
        record_id = uuid.uuid4().hex[:12]

        # Store in pending records
        pending_records[record_id] = {
            "sales_name": sales_name,
            "customer": customer,
            "transcript": full_transcript,
            "context": context,
            "extracted": extracted,
            "provider_used": result.provider_used,
            "fallback_used": result.fallback_used,
            "elapsed_seconds": result.elapsed_seconds,
            "confidence": result.confidence,
            "passed": result.passed,
            "confirmed": False,
        }

        # ---- Confidence-based auto-confirmation ----
        # Average only extraction-reliant fields that have non-zero confidence
        field_conf = result.confidence or {}
        extract_fields = ["contact_person","visit_time","topic","key_concerns",
                         "competitors_mentioned","next_steps","risk_level","risk_reason"]
        extract_scores = [float(v) for f in extract_fields for k,v in field_conf.items() 
                         if k == f and float(v) > 0]
        if not extract_scores:
            extract_scores = [float(v) for v in field_conf.values() if float(v) > 0]
        overall_conf = sum(extract_scores) / len(extract_scores) if extract_scores else 0

        # Additional safety checks
        auto_confirmed = False
        transcript_len = len(transcript)
        has_substance = transcript_len >= 25  # minimum meaningful length

        if overall_conf > 0.85 and result.passed and has_substance:
            # High confidence + all checks passed → auto-confirm
            auto_confirmed = True
            pending_records[record_id]["confirmed"] = True
            try:
                _handle_confirm(pending_records[record_id], record_id)
            except Exception as e:
                print(f"[Auto-confirm error] {e}")
                auto_confirmed = False
                pending_records[record_id]["confirmed"] = False

        # Send Feishu notification
        try:
            notifier.send_confirmation_card(extracted, record_id, auto_confirmed=auto_confirmed)
        except Exception as e:
            print(f"[Feishu notify failed] {e}")

        return jsonify({
            "success": True,
            "record_id": record_id,
            "auto_confirmed": auto_confirmed,
            "confidence": round(overall_conf, 2),
        })

    except Exception as e:
        print(f"[API analyze error] {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/record/<record_id>", methods=["GET"])
def api_get_record(record_id: str):
    """Get pending record by ID."""
    if record_id not in pending_records:
        return jsonify({"success": False, "error": "记录不存在"}), 404

    rec = pending_records[record_id]
    return jsonify({
        "success": True,
        "extracted": rec["extracted"],
        "transcript": rec["transcript"],
        "provider_used": rec["provider_used"],
        "fallback_used": rec["fallback_used"],
        "elapsed_seconds": rec["elapsed_seconds"],
        "confidence": rec["confidence"],
    })


@app.route("/api/stats", methods=["GET"])
def api_stats():
    """Get system statistics."""
    records = _load_all_records()
    customers = set()
    for r in records:
        c = r.get("Customer") or r.get("customer") or ""
        if c:
            customers.add(c)
    pending_count = sum(1 for v in pending_records.values() if not v.get("confirmed"))

    return jsonify({
        "total": len(records),
        "pending": pending_count,
        "customers": len(customers),
    })


@app.route("/api/record/<record_id>/action", methods=["POST"])
def api_record_action(record_id: str):
    """Handle confirm/modify/cancel actions on a pending record."""
    if record_id not in pending_records:
        return jsonify({"success": False, "error": "记录不存在或已过期"}), 404

    rec = pending_records[record_id]
    action = request.json.get("action", "").lower()

    if action == "confirm":
        return _handle_confirm(rec, record_id)

    elif action == "modify":
        modifications = request.json.get("modifications", {})
        return _handle_modify(rec, record_id, modifications)

    elif action == "cancel":
        del pending_records[record_id]
        return jsonify({"success": True, "message": "记录已取消"})

    else:
        return jsonify({"success": False, "error": f"未知操作: {action}"}), 400


def _handle_confirm(rec: dict, record_id: str) -> dict:
    """Handle confirm action: save to CSV, log, update memory, notify Feishu."""
    extracted = rec["extracted"]
    transcript = rec["transcript"]

    # Save to CSV
    row_id = _save_to_csv(extracted, transcript)

    # Write audit log
    try:
        logger = ExcelAuditLogger()
        entry = AuditEntry(
            transcript=transcript,
            context=rec.get("context", {}),
            extracted=extracted,
            provider_used=rec.get("provider_used", ""),
            fallback_used=rec.get("fallback_used", False),
            passed=rec.get("passed", False),
            human_review_required=False,
            human_reason="",
            errors=[],
            introspection_rounds=0,
            confidence=rec.get("confidence", {}),
            elapsed_seconds=rec.get("elapsed_seconds", 0.0),
            feishu_status="已确认入库",
        )
        logger.log(entry)
    except Exception as e:
        print(f"[Audit log failed] {e}")

    # Update memory store
    try:
        memory_store.update_with_extraction(
            customer_name=extracted.get("customer", ""),
            extracted_data=extracted,
            transcript=transcript,
        )
    except Exception as e:
        print(f"[Memory update failed] {e}")

    # Send Feishu confirmation notification
    try:
        customer = extracted.get("customer", "")
        sales = extracted.get("sales_name", "")
        risk = extracted.get("risk_level", "")
        notifier.send_text(f"销售拜访记录已确认入库\n销售：{sales} | 客户：{customer} | 风险等级：{risk} | 记录ID：{row_id}")
    except Exception as e:
        print(f"[Feishu confirm notify failed] {e}")

    # Mark as confirmed
    rec["confirmed"] = True

    return jsonify({"success": True, "row_id": row_id})


def _handle_modify(rec: dict, record_id: str, modifications: dict) -> dict:
    """Handle modify action: apply changes then save."""
    # Apply modifications to extracted data
    extracted = rec["extracted"].copy()
    extracted.update(modifications)
    rec["extracted"] = extracted
    rec["transcript"] = rec.get("transcript", "")

    # Then confirm (save to CSV)
    return _handle_confirm(rec, record_id)


# ============================================================================
# Main
# ============================================================================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print("=" * 60)
    print("  销售拜访 AI 分析系统 - Web 版")
    print("=" * 60)
    print(f"  启动地址: http://localhost:{port}")
    print("  按 Ctrl+C 停止")
    print("=" * 60)
    app.run(host="127.0.0.1", port=port, debug=False)
